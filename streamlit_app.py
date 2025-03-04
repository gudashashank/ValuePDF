import streamlit as st
import openai
from PyPDF2 import PdfReader
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
import tiktoken
import io
import httpx
from langdetect import detect

# Custom HTTP Client to handle proxies argument
class CustomHTTPClient(httpx.Client):
    def __init__(self, *args, **kwargs):
        kwargs.pop("proxies", None)
        super().__init__(*args, **kwargs)

# Streamlit App Configuration
st.set_page_config(
    page_title="SmartDoc Converter",
    page_icon="📝",
    layout="wide"
)

# Initialize session state
if 'api_key' not in st.session_state:
    st.session_state.api_key = None
if 'authenticated' not in st.session_state:
    st.session_state.authenticated = False
if 'openai_client' not in st.session_state:
    st.session_state.openai_client = None

# Function to count tokens
def count_tokens(text, model="gpt-4o-mini"):
    encoding = tiktoken.encoding_for_model(model)
    return len(encoding.encode(text))

# Agent-like function to call GPT-4o-mini
def call_gpt_agent(prompt, client, max_tokens=4096, temperature=0.7):
    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert assistant with specialized skills."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=max_tokens,
            temperature=temperature
        )
        return response.choices[0].message.content, response.usage.total_tokens
    except Exception as e:
        st.error(f"Agent error: {str(e)}")
        return None, 0

# Function to extract text and estimate font sizes from PDF
def extract_text_and_fonts_from_pdf(pdf_file, client):
    try:
        reader = PdfReader(pdf_file)
        text = ""
        font_sizes = []
        for page in reader.pages:
            text += page.extract_text()
            if "/Resources" in page and "/Font" in page["/Resources"]:
                font_dict = page["/Resources"]["/Font"]
                if font_dict:  # Ensure font_dict is not None
                    for font_ref in font_dict.values():
                        # Resolve IndirectObject to get the actual font dictionary
                        font = font_ref.get_object() if hasattr(font_ref, 'get_object') else font_ref
                        if isinstance(font, dict) and "/Size" in font:
                            font_sizes.append(font["/Size"])
        
        # If no font sizes detected, use GPT-4o-mini to estimate
        if not font_sizes:
            prompt = f"Estimate the average font size (in points) of this text based on typical document styles:\n\n{text[:1000]}"
            estimated_size, tokens = call_gpt_agent(prompt, client, max_tokens=50, temperature=0.5)
            avg_font_size = float(estimated_size.strip()) if estimated_size and estimated_size.strip().replace('.','',1).isdigit() else 12
        else:
            avg_font_size = sum(font_sizes) / len(font_sizes)
        return text, avg_font_size
    except Exception as e:
        st.error(f"Error extracting PDF text: {str(e)}")
        return "", 12

# Function to convert text to Word document with font size and highlighting
def text_to_word(text, font_size=12, highlight_terms=None):
    doc = Document()
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = int(font_size * 13333)  # 1 pt = 13333 in Word
    
    if highlight_terms:
        paragraph.clear()
        for i, part in enumerate(text.split(highlight_terms[0])):  # Simplified for first term
            paragraph.add_run(part)
            if i < len(text.split(highlight_terms[0])) - 1:
                highlighted_run = paragraph.add_run(highlight_terms[0])
                highlighted_run.font.highlight_color = 7  # Yellow
                highlighted_run.font.size = int(font_size * 13333)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Agent for filename suggestion
def suggest_filename(text, client):
    prompt = f"Suggest a concise, meaningful filename (without extension) based on this text:\n\n{text[:500]}"
    filename, tokens = call_gpt_agent(prompt, client, max_tokens=20, temperature=0.3)
    return filename.strip() if filename else "converted_output"

# Agent for key term identification
def identify_key_terms(text, client):
    prompt = f"Identify 3-5 key terms or phrases from this text that should be highlighted:\n\n{text[:1000]}"
    terms, tokens = call_gpt_agent(prompt, client, max_tokens=50, temperature=0.5)
    return [term.strip() for term in terms.split(",")] if terms else []

# Agent for summarization
def summarize_text(text, client):
    prompt = f"Provide a concise summary (100-150 words) of this text:\n\n{text}"
    summary, tokens = call_gpt_agent(prompt, client, max_tokens=200, temperature=0.7)
    return summary, tokens

# Agent for formatting optimization and conversion
def convert_and_optimize_format(pdf_text, target_format="word", client=None, translate_to=None, highlight_terms=None):
    prompt = f"Convert this text into {target_format} format with optimized structure (e.g., headings, bullet points where appropriate):\n\n{pdf_text}"
    if translate_to:
        prompt += f"\nTranslate the output to {translate_to}."
    if highlight_terms:
        prompt += f"\nIndicate where to highlight these terms: {', '.join(highlight_terms)}."
    
    formatted_text, total_tokens = call_gpt_agent(prompt, client)
    return formatted_text, total_tokens

# Streamlit App
st.title("SmartDoc Converter with AI Agents 📝")
st.subheader("Convert PDFs with intelligent agents for formatting, translation, and more!")

# Step 1: Input OpenAI API Key
api_key = st.text_input("Enter your OpenAI API Key", type="password", value=st.session_state.api_key if st.session_state.api_key else "")
if api_key and not st.session_state.authenticated:
    try:
        client = openai.OpenAI(api_key=api_key, http_client=CustomHTTPClient())
        client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": "test"}],
            max_tokens=5
        )
        st.session_state.api_key = api_key
        st.session_state.openai_client = client
        st.session_state.authenticated = True
        st.success("API Key Authenticated!")
    except openai.AuthenticationError:
        st.error("Invalid API Key. Please check and try again.")
    except Exception as e:
        st.error(f"Authentication error: {str(e)}")

# Step 2: Upload PDF and Process
if st.session_state.authenticated:
    uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])
    if uploaded_file:
        st.info("Processing your PDF with AI agents...")

        # Extract text and font size
        with st.spinner("Agent 1: Extracting text and estimating font size..."):
            pdf_text, avg_font_size = extract_text_and_fonts_from_pdf(uploaded_file, st.session_state.openai_client)
            if pdf_text:
                st.text_area("Preview of Extracted Text", pdf_text, height=200)
                st.write(f"Detected/Estimated Font Size: {avg_font_size:.1f} pt")
            else:
                st.warning("No text extracted from PDF.")

        if pdf_text:
            # UI for options
            col1, col2 = st.columns(2)
            with col1:
                target_format = st.selectbox("Output Format", ["Word (.docx)", "Markdown (.md)"])
                suggested_filename = suggest_filename(pdf_text, st.session_state.openai_client)
                output_filename = st.text_input("Output Filename (without extension)", value=suggested_filename)

            with col2:
                detected_lang = detect(pdf_text)
                st.write(f"Detected Language: {detected_lang}")
                translate_to = st.selectbox("Translate to (optional)", ["None", "English", "Spanish", "French", "German"], index=0)
                translate_to = None if translate_to == "None" else translate_to

                auto_terms = identify_key_terms(pdf_text, st.session_state.openai_client)
                st.write(f"Suggested Terms to Highlight: {', '.join(auto_terms)}")
                highlight_input = st.text_input("Highlight Terms (comma-separated, Word only)", "")
                highlight_terms = [term.strip() for term in highlight_input.split(",")] if highlight_input else auto_terms if target_format == "Word (.docx)" else None

            add_summary = st.checkbox("Add Summary to Output")

            if st.button("Convert PDF"):
                with st.spinner("Agents at work..."):
                    status = st.empty()
                    total_tokens_used = 0

                    if add_summary:
                        status.write("Agent 2: Generating summary...")
                        summary, summary_tokens = summarize_text(pdf_text, st.session_state.openai_client)
                        total_tokens_used += summary_tokens
                        if summary:
                            pdf_text = f"Summary:\n{summary}\n\nOriginal Text:\n{pdf_text}"

                    status.write("Agent 3: Optimizing and converting format...")
                    formatted_text, conversion_tokens = convert_and_optimize_format(
                        pdf_text, 
                        target_format.split()[0].lower(), 
                        st.session_state.openai_client,
                        translate_to,
                        highlight_terms
                    )
                    total_tokens_used += conversion_tokens

                    if not formatted_text:
                        raise ValueError("Conversion returned no text.")

                    status.write("Agent 4: Preparing output file...")
                    if target_format == "Word (.docx)":
                        output_filename_full = f"{output_filename}.docx"
                        buffer = text_to_word(formatted_text, font_size=avg_font_size, highlight_terms=highlight_terms)
                        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        download_content = buffer.getvalue()
                    elif target_format == "Markdown (.md)":
                        output_filename_full = f"{output_filename}.md"
                        buffer = io.BytesIO(formatted_text.encode('utf-8'))
                        mime_type = "text/markdown"
                        download_content = buffer.getvalue()

                    status.write("Conversion complete!")
                    st.success("All agents completed their tasks!")

                    cost_per_1k_tokens = 0.015  # Hypothetical cost
                    cost = (total_tokens_used / 1000) * cost_per_1k_tokens
                    st.info(f"Total Tokens Used: {total_tokens_used}")
                    st.info(f"Estimated Cost: ${cost:.4f}")

                    st.download_button(
                        label=f"Download Converted File ({target_format})",
                        data=download_content,
                        file_name=output_filename_full,
                        mime=mime_type
                    )

                    status.empty()

else:
    st.warning("Please enter and authenticate your API key to proceed.")

# Reset API Key option
if st.session_state.authenticated and st.button("Reset API Key"):
    st.session_state.api_key = None
    st.session_state.openai_client = None
    st.session_state.authenticated = False
    st.rerun()